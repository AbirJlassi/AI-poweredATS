"""
CV Parser Classique v3 — Regex + SpaCy NER  (FR + EN)
=======================================================
  - Détection automatique de la langue (FR / EN) via heuristique de mots-clés
  - Chargement dynamique du modèle SpaCy selon la langue détectée
  - SECTION_HEADERS : couverture complète FR + EN
      (work experience, professional background, achievements, awards,
       qualifications, career, volunteer, references…)
  - DATE_RANGE : mois anglais complets et abrégés (January / Jan)
  - LanguagesExtractor : niveaux EN (native, fluent, proficient, B2, C1…)
  - CertificationsExtractor : keywords EN (honor, award, achievement, credential)
  - SectionSplitter._KEY_MAP : clés EN normalisées vers les clés canoniques
  - Tous les _find_section() : listes de clés FR + EN

Dépendances :
    pip install spacy pdfminer.six python-docx
    python -m spacy download fr_core_news_lg
    python -m spacy download en_core_web_lg
"""

import re
import json
import unicodedata
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional

import spacy
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document


# ---------------------------------------------------------------------------
# Data Model
# ---------------------------------------------------------------------------

@dataclass
class CVData:
    full_name:      Optional[str] = None
    email:          Optional[str] = None
    phone:          Optional[str] = None
    location:       Optional[str] = None
    linkedin:       Optional[str] = None
    github:         Optional[str] = None
    summary:        Optional[str] = None
    skills:         list[str]     = field(default_factory=list)
    languages:      list[str]     = field(default_factory=list)
    experiences:    list[dict]    = field(default_factory=list)
    education:      list[dict]    = field(default_factory=list)
    certifications: list[str]     = field(default_factory=list)
    detected_lang:  Optional[str] = None   # "fr" | "en" | "unknown"
    _confidence:    dict          = field(default_factory=dict)

    def to_json(self, indent: int = 2) -> str:
        d = asdict(self)
        d.pop("_confidence", None)
        return json.dumps(d, ensure_ascii=False, indent=indent)

    def to_json_with_confidence(self, indent: int = 2) -> str:
        return json.dumps(asdict(self), ensure_ascii=False, indent=indent)


# ---------------------------------------------------------------------------
# Language Detector  ← NOUVEAU v3
# ---------------------------------------------------------------------------

class LanguageDetector:
    """
    Détecte FR ou EN par comptage de mots-clés caractéristiques.
    Léger, sans dépendance externe (pas de langdetect).
    """

    _FR_KEYWORDS = re.compile(
        r"\b(expérience|formation|compétences|langues|diplôme|poste|"
        r"entreprise|société|mois|année|depuis|jusqu|présent|"
        r"responsabilités|missions|réalisations|secteur)\b",
        re.IGNORECASE,
    )
    _EN_KEYWORDS = re.compile(
        r"\b(experience|education|skills|languages|degree|position|"
        r"company|month|year|since|until|present|current|"
        r"responsibilities|achievements|industry|summary|objective)\b",
        re.IGNORECASE,
    )

    @classmethod
    def detect(cls, text: str) -> str:
        sample = text[:3000]
        fr_count = len(cls._FR_KEYWORDS.findall(sample))
        en_count = len(cls._EN_KEYWORDS.findall(sample))
        if fr_count == 0 and en_count == 0:
            return "unknown"
        return "fr" if fr_count >= en_count else "en"


# ---------------------------------------------------------------------------
# Text Normalization
# ---------------------------------------------------------------------------

class TextNormalizer:
    _LIGATURES = str.maketrans({
        "\ufb00": "ff", "\ufb01": "fi", "\ufb02": "fl",
        "\ufb03": "ffi", "\ufb04": "ffl",
        "\u2019": "'",
        "\u00a0": " ",
        "\u200b": "",
    })

    @classmethod
    def normalize(cls, text: str) -> str:
        text = unicodedata.normalize("NFKC", text)
        text = text.translate(cls._LIGATURES)
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[–—]", "-", text)
        text = re.sub(r"[^\S\n]+", " ", text)
        return text.strip()


# ---------------------------------------------------------------------------
# Text Extraction
# ---------------------------------------------------------------------------

class TextExtractor:
    @staticmethod
    def extract(file_path: str) -> str:
        path   = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            raw = pdf_extract_text(file_path) or ""
        elif suffix == ".docx":
            doc   = Document(file_path)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            raw = "\n".join(parts)
        elif suffix in (".txt", ".md"):
            raw = path.read_text(encoding="utf-8")
        else:
            raise ValueError(f"Format non supporté : {suffix}")

        return TextNormalizer.normalize(raw)


# ---------------------------------------------------------------------------
# Regex Patterns  ← DATE_RANGE étendu EN v3
# ---------------------------------------------------------------------------

class RegexPatterns:
    EMAIL = re.compile(
        r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
    )
    PHONE = re.compile(
        r"(?:\+\d{1,3}[\s.\-]?)?(?:\(?\d{2,4}\)?[\s.\-]?){2,5}\d{2,4}"
    )
    LINKEDIN = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+")
    GITHUB   = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w\-]+")

    # ------------------------------------------------------------------ #
    # SECTION HEADERS — FR + EN complet  ← v3
    # ------------------------------------------------------------------ #
    SECTION_HEADERS = re.compile(
        r"^\s*[-=*#\u2022]*\s*("

        # ── Expériences / Experience ──────────────────────────────────
        r"exp[eé]riences?\s*(professionnelles?|de\s+travail)?|"
        r"parcours\s*professionnel|"
        r"work\s*experience|professional\s*(experience|background|history)|"
        r"employment\s*(history|record)?|career\s*(history|summary)?|"
        r"work\s*history|positions?\s*held|"

        # ── Formation / Education ─────────────────────────────────────
        r"formations?\s*(acad[eé]miques?|universitaires?)?|"
        r"[eé]ducation|dipl[oô]mes?|"
        r"academic\s*(background|qualifications?|history)?|"
        r"qualifications?|degrees?|schooling|"

        # ── Compétences / Skills ──────────────────────────────────────
        r"comp[eé]tences?\s*(techniques?|cl[eé]s?|professionnelles?)?|"
        r"skills?\s*(summary|set)?|technical\s*skills?|"
        r"core\s*(skills?|competenc(ies|es))|"
        r"technologies|outils?|tools?|"
        r"areas?\s*of\s*(expertise|strength)|expertise|"

        # ── Langues / Languages ───────────────────────────────────────
        r"langues?\s*(ma[iî]tris[eé]es?|parl[eé]es?)?|"
        r"languages?\s*(spoken|known)?|"

        # ── Certifications ────────────────────────────────────────────
        r"certifications?|attestations?|licen[sc]es?|credentials?|"
        r"professional\s*development|training|"

        # ── Projets / Projects ────────────────────────────────────────
        r"projets?\s*(personnels?|professionnels?|acad[eé]miques?)?|"
        r"projects?\s*(portfolio)?|portfolio|"
        r"personal\s*projects?|side\s*projects?|"

        # ── Résumé / Summary ──────────────────────────────────────────
        r"r[eé]sum[eé]|profil\s*(professionnel)?|"
        r"summary|professional\s*summary|"
        r"objective|career\s*objective|"
        r"about(\s*me)?|profile|"

        # ── Récompenses / Awards ──────────────────────────────────────
        r"r[eé]compenses?|distinctions?|prix|"
        r"awards?|honors?|honours?|achievements?|recognitions?|"

        # ── Bénévolat / Volunteer ─────────────────────────────────────
        r"b[eé]n[eé]volat|associations?|activit[eé]s?\s*extra.*|"
        r"volunteer(ing)?|extracurricular|civic\s*activities|"

        # ── Références / References ───────────────────────────────────
        r"r[eé]f[eé]rences?|references?|"

        # ── Divers ────────────────────────────────────────────────────
        r"loisirs?|int[eé]r[eê]ts?|hobbies?|interests?"

        r")\s*[-=:*#\u2022]*\s*$",
        re.IGNORECASE | re.MULTILINE,
    )

    # ------------------------------------------------------------------ #
    # DATE_RANGE — FR + EN  ← v3
    # ------------------------------------------------------------------ #
    _MONTHS_FR = (
        r"jan(?:vier)?|f[eé]v(?:rier)?|mar(?:s)?|avr(?:il)?|mai|juin|"
        r"juil(?:let)?|ao[uû]t|sep(?:t(?:embre)?)?|oct(?:obre)?|"
        r"nov(?:embre)?|d[eé]c(?:embre)?"
    )
    _MONTHS_EN = (
        r"jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
        r"jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|"
        r"nov(?:ember)?|dec(?:ember)?"
    )
    _MONTHS    = f"(?:{_MONTHS_FR}|{_MONTHS_EN})"
    _PRESENT   = r"aujourd'hui|pr[eé]sent|present|en\s+cours|maintenant|current(?:ly)?|now|ongoing|to\s+date"
    _DATE_PART = f"(?:{_MONTHS}\\.?\\s*\\d{{4}}|\\d{{2}}/\\d{{4}}|\\d{{4}})"

    DATE_RANGE = re.compile(
        rf"(?:{_DATE_PART})\s*[-/]\s*(?:{_DATE_PART}|{_PRESENT})",
        re.IGNORECASE,
    )


# ---------------------------------------------------------------------------
# SpaCy NER — chargement dynamique selon la langue  ← v3
# ---------------------------------------------------------------------------

_SPACY_MODEL_MAP = {
    "fr":      "fr_core_news_lg",
    "en":      "en_core_web_lg",
    "unknown": "fr_core_news_lg",   # fallback
}

class NERExtractor:
    """Charge le modèle SpaCy adapté à la langue détectée."""

    def __init__(self, lang: str = "fr"):
        model = _SPACY_MODEL_MAP.get(lang, "fr_core_news_lg")
        try:
            self.nlp = spacy.load(model)
        except OSError:
            # Essai avec le modèle sm si lg absent
            model_sm = model.replace("_lg", "_sm").replace("_md", "_sm")
            try:
                self.nlp = spacy.load(model_sm)
            except OSError:
                raise OSError(
                    f"Modèle SpaCy '{model}' (ni '{model_sm}') introuvable.\n"
                    f"Installez : python -m spacy download {model}"
                )

    def extract_entities(self, text: str) -> dict[str, list[str]]:
        doc      = self.nlp(text[:5000])
        entities: dict[str, list[str]] = {}
        for ent in doc.ents:
            val = ent.text.strip()
            if val and val not in entities.get(ent.label_, []):
                entities.setdefault(ent.label_, []).append(val)
        return entities

    def get_all(self, text: str) -> dict:
        ents = self.extract_entities(text)
        return {
            "name":      ents.get("PER", [None])[0],
            "locations": ents.get("LOC", []) + ents.get("GPE", []),
            "orgs":      ents.get("ORG", []),
        }


# ---------------------------------------------------------------------------
# Section Splitter  ← _KEY_MAP étendu
# ---------------------------------------------------------------------------

class SectionSplitter:

    _KEY_MAP = {
        # ── FR ──
        "expériences": "experiences",
        "expériences professionnelles": "experiences",
        "experience professionnelle": "experiences",
        "parcours professionnel": "experiences",
        "formations": "education",
        "formation": "education",
        "formation académique": "education",
        "éducation": "education",
        "education": "education",
        "diplômes": "education",
        "compétences": "skills",
        "compétences techniques": "skills",
        "compétences clés": "skills",
        "technologies": "skills",
        "outils": "skills",
        "langues": "languages",
        "langues maîtrisées": "languages",
        "certifications": "certifications",
        "attestations": "certifications",
        "projets": "projects",
        "projets personnels": "projects",
        "résumé": "summary",
        "profil": "summary",
        "profil professionnel": "summary",
        # ── EN ──
        "work experience": "experiences",
        "professional experience": "experiences",
        "professional background": "experiences",
        "employment history": "experiences",
        "employment record": "experiences",
        "career history": "experiences",
        "career summary": "experiences",
        "work history": "experiences",
        "positions held": "experiences",
        "academic background": "education",
        "academic qualifications": "education",
        "academic history": "education",
        "qualifications": "education",
        "degrees": "education",
        "schooling": "education",
        "skills": "skills",
        "skills summary": "skills",
        "technical skills": "skills",
        "core skills": "skills",
        "core competencies": "skills",
        "areas of expertise": "skills",
        "expertise": "skills",
        "tools": "skills",
        "languages": "languages",
        "languages spoken": "languages",
        "licenses": "certifications",
        "credentials": "certifications",
        "professional development": "certifications",
        "training": "certifications",
        "projects": "projects",
        "personal projects": "projects",
        "side projects": "projects",
        "portfolio": "projects",
        "summary": "summary",
        "professional summary": "summary",
        "objective": "summary",
        "career objective": "summary",
        "about": "summary",
        "about me": "summary",
        "profile": "summary",
        "awards": "awards",
        "honors": "awards",
        "achievements": "awards",
        "recognitions": "awards",
        "distinctions": "awards",
        "volunteer": "volunteer",
        "volunteering": "volunteer",
        "extracurricular": "volunteer",
        "references": "references",
        "interests": "interests",
        "hobbies": "interests",
    }

    # Mots-clés pour le fallback MAJUSCULES
    _UPPER_HINTS = [
        "exp", "work", "employ", "career",
        "form", "educ", "school", "degree", "qualif",
        "comp", "skill", "tech", "tool", "expert",
        "lang",
        "cert", "licen", "train",
        "proj", "portfo",
        "profil", "summar", "object", "about",
        "award", "honor", "achiev",
        "volunt",
    ]

    def _normalize_key(self, raw: str) -> str:
        clean = re.sub(r"[-=:*#\u2022]", "", raw).strip().lower()
        clean = re.sub(r"\s+", " ", clean)
        return self._KEY_MAP.get(clean, clean)

    def split(self, text: str) -> dict[str, str]:
        sections: dict[str, str] = {}
        current_section = "header"
        current_content: list[str] = []

        for line in text.splitlines():
            is_header = bool(RegexPatterns.SECTION_HEADERS.match(line))

            if not is_header:
                stripped = line.strip()
                if (len(stripped) >= 4
                        and stripped.isupper()
                        and not re.match(r"^\d", stripped)):
                    is_header = any(
                        hint in stripped.lower()
                        for hint in self._UPPER_HINTS
                    )

            if is_header:
                sections[current_section] = "\n".join(current_content).strip()
                current_section = self._normalize_key(line)
                current_content = []
            else:
                current_content.append(line)

        sections[current_section] = "\n".join(current_content).strip()
        return sections


# ---------------------------------------------------------------------------
# Skills Extractor
# ---------------------------------------------------------------------------

TECH_SKILLS_DICTIONARY = {
    # Langages
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
    "php", "ruby", "scala", "kotlin", "swift", "r", "matlab", "bash", "sql",
    "perl", "lua", "dart", "haskell", "elixir",
    # Web
    "react", "vue", "angular", "next.js", "nuxt", "django", "flask",
    "fastapi", "spring", "node.js", "express", "html", "css", "sass",
    "graphql", "rest", "tailwind", "bootstrap",
    # Data / ML / AI
    "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy", "spark",
    "kafka", "airflow", "dbt", "mlflow", "opencv", "nltk", "hugging face",
    "transformers", "xgboost", "lightgbm", "keras", "spacy",
    # Cloud & DevOps
    "aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ansible",
    "ci/cd", "github actions", "jenkins", "gitlab ci", "helm",
    # Bases de données
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
    "cassandra", "snowflake", "bigquery", "sqlite", "oracle", "mariadb",
    # Outils
    "git", "linux", "jira", "figma", "powerbi", "tableau", "excel",
    "latex", "jupyter", "vscode", "postman", "swagger",
}

class SkillsExtractor:
    def extract(self, text: str) -> list[str]:
        text_lower = text.lower()
        found = set()
        for skill in TECH_SKILLS_DICTIONARY:
            pattern = r"(?<![a-z])" + re.escape(skill) + r"(?![a-z])"
            if re.search(pattern, text_lower):
                found.add(skill)
        return sorted(found)


# ---------------------------------------------------------------------------
# Languages Extractor  ← niveaux EN complets v3
# ---------------------------------------------------------------------------

KNOWN_LANGUAGES = {
    # FR labels
    "arabe": "Arabic", "arabic": "Arabic",
    "français": "French", "french": "French",
    "anglais": "English", "english": "English",
    "espagnol": "Spanish", "spanish": "Spanish",
    "allemand": "German", "german": "German",
    "italien": "Italian", "italian": "Italian",
    "chinois": "Chinese", "chinese": "Chinese",
    "japonais": "Japanese", "japanese": "Japanese",
    "russe": "Russian", "russian": "Russian",
    "portugais": "Portuguese", "portuguese": "Portuguese",
    "turc": "Turkish", "turkish": "Turkish",
    "néerlandais": "Dutch", "dutch": "Dutch",
    "coréen": "Korean", "korean": "Korean",
    "hébreu": "Hebrew", "hebrew": "Hebrew",
    "persan": "Persian", "persian": "Persian",
    "hindi": "Hindi",
}

PROFICIENCY_LEVELS = re.compile(
    r"("
    # CECRL
    r"[abcABC][12]"
    r"|"
    # FR
    r"natif(?:ve)?|langue\s+maternelle|courant|bilingue|"
    r"avanc[eé]|interm[eé]diaire|notions?|d[eé]butant"
    r"|"
    # EN
    r"native|mother\s*tongue|fluent|bilingual|"
    r"proficient|advanced|upper[\s-]intermediate|intermediate|"
    r"elementary|basic|beginner|working\s*proficiency|"
    r"limited\s*working|full\s*professional"
    r")",
    re.IGNORECASE,
)

class LanguagesExtractor:
    def extract(self, section_text: str) -> list[str]:
        results = []
        text_lower = section_text.lower()

        for keyword, canonical in KNOWN_LANGUAGES.items():
            if re.search(r"\b" + re.escape(keyword) + r"\b", text_lower):
                idx     = text_lower.find(keyword)
                context = section_text[idx: idx + 80]
                level_m = PROFICIENCY_LEVELS.search(context)
                entry   = f"{canonical} ({level_m.group(0).lower()})" if level_m else canonical
                results.append((canonical, entry))

        # Déduplique par langue canonique
        seen: set[str] = set()
        deduped = []
        for canonical, entry in results:
            if canonical not in seen:
                seen.add(canonical)
                deduped.append(entry)
        return deduped


# ---------------------------------------------------------------------------
# Certifications Extractor  ← keywords EN ajoutés v3
# ---------------------------------------------------------------------------

CERT_KEYWORDS = re.compile(
    r"(?:"
    # FR
    r"certification|certifi[eé]|attestation|diplôme|licence|badge|"
    # EN
    r"certified|certificate|credential|accreditation|"
    r"honor|honour|award|achievement|recognition|"
    # Vendors
    r"aws\s+certified|google\s+cloud|microsoft\s+certified|"
    r"pmp|scrum|itil|cisco|oracle\s+certified|comptia|"
    r"cka|ckad|cks|rhce|rhcsa|lpic|gcp\s+associate"
    r")",
    re.IGNORECASE,
)

class CertificationsExtractor:
    def extract(self, section_text: str) -> list[str]:
        certs = []
        for line in section_text.splitlines():
            line = line.strip()
            if line and CERT_KEYWORDS.search(line):
                clean = re.sub(r"^[-•*►▪\u2022]\s*", "", line)
                if len(clean) > 5:
                    certs.append(clean)
        return certs


# ---------------------------------------------------------------------------
# Experience & Education Parsers
# ---------------------------------------------------------------------------

class ExperienceParser:

    def __init__(self, ner: NERExtractor):
        self.ner = ner

    def parse(self, section_text: str) -> list[dict]:

        experiences = []
        lines = [l.strip() for l in section_text.splitlines() if l.strip()]

        for line in lines:

            date_match = RegexPatterns.DATE_RANGE.search(line)

            if not date_match:
                continue

            date_range = date_match.group(0)

            before_date = line[:date_match.start()].strip(" |-,")

            ents = self.ner.get_all(before_date)

            company = ents["orgs"][0] if ents["orgs"] else None

            position = before_date
            if company:
                position = position.replace(company, "").strip(" |-,")

            experiences.append({
                "position": position if position else None,
                "company": company,
                "date_range": date_range
            })

        return experiences


class EducationParser:

    def __init__(self, ner: NERExtractor):
        self.ner = ner

    def parse(self, section_text: str) -> list[dict]:

        education = []
        lines = [l.strip() for l in section_text.splitlines() if l.strip()]

        for line in lines:

            date_match = RegexPatterns.DATE_RANGE.search(line)

            date_range = date_match.group(0) if date_match else None

            if date_match:
                before_date = line[:date_match.start()].strip(" |-,")
            else:
                before_date = line

            ents = self.ner.get_all(before_date)

            school = ents["orgs"][0] if ents["orgs"] else None

            degree = before_date
            if school:
                degree = degree.replace(school, "").strip(" |-,")

            education.append({
                "degree": degree if degree else None,
                "school": school,
                "date_range": date_range
            })

        return education


# ---------------------------------------------------------------------------
# Main CV Parser
# ---------------------------------------------------------------------------

class ClassicCVParser:
    """
    Parser FR/EN v3 : détecte la langue, charge le bon modèle SpaCy,
    applique des patterns et dictionnaires bilingues.
    """

    def __init__(self, spacy_model: Optional[str] = None):
        """
        spacy_model : forcer un modèle SpaCy précis (ex: "en_core_web_lg").
                      Si None, le modèle est choisi automatiquement selon la
                      langue détectée dans le document.
        """
        self._forced_model = spacy_model
        self._ner_cache: dict[str, NERExtractor] = {}
        

        self.splitter   = SectionSplitter()
        self.skills_ext = SkillsExtractor()
        self.lang_ext   = LanguagesExtractor()
        self.cert_ext   = CertificationsExtractor()
        # parsers created once we know which SpaCy model/NER extractor to use
        self.exp_parser: Optional[ExperienceParser] = None
        self.edu_parser: Optional[EducationParser] = None

    def _get_ner(self, lang: str) -> NERExtractor:
        """Charge (et met en cache) le NERExtractor pour une langue."""
        if lang not in self._ner_cache:
            if self._forced_model:
                extractor = NERExtractor.__new__(NERExtractor)
                extractor.nlp = spacy.load(self._forced_model)
            else:
                extractor = NERExtractor(lang)
            self._ner_cache[lang] = extractor
        return self._ner_cache[lang]

    def parse(self, file_path: str) -> CVData:
        # 1. Extraction + normalisation
        raw_text = TextExtractor.extract(file_path)

        # 2. Détection de langue
        lang = LanguageDetector.detect(raw_text)
        ner  = self._get_ner(lang)

        # instancie les parsers dépendant du ner (récréation possible selon la langue)
        self.exp_parser = ExperienceParser(ner)
        self.edu_parser = EducationParser(ner)

        # 3. Découpage en sections
        sections = self.splitter.split(raw_text)
        header   = sections.get("header", raw_text[:1500])

        # DEBUG : décommenter pour inspecter les sections détectées
        #print(f"[v3] Langue: {lang} | Sections: {list(sections.keys())}")

        # 4. Contact (Regex — language-agnostic)
        email_m    = RegexPatterns.EMAIL.search(raw_text)
        phone_m    = RegexPatterns.PHONE.search(raw_text)
        linkedin_m = RegexPatterns.LINKEDIN.search(raw_text)
        github_m   = RegexPatterns.GITHUB.search(raw_text)

        # 5. NER
        ner_header = ner.get_all(header)
        ner_full   = ner.get_all(raw_text)

        # 6. Compétences
        skills_section = self._find_section(
            sections,
            ["skills", "compétences", "technical skills", "core skills",
             "technologies", "outils", "tools", "expertise", "areas of expertise"]
        )
        skills = self.skills_ext.extract(skills_section or raw_text)

        # 7. Langues
        lang_section = self._find_section(
            sections,
            ["languages", "langues", "languages spoken", "langues maîtrisées"]
        )
        languages = self.lang_ext.extract(lang_section or raw_text)

        # 8. Expériences
        exp_section = self._find_section(
            sections,
            ["experiences","expériences", "work experience", "professional experience",
             "employment history", "career history", "work history",
             "parcours", "expériences professionnelles", "positions held"]
        )
        experiences = self.exp_parser.parse(exp_section) if exp_section else []

        # 9. Formation
        edu_section = self._find_section(
            sections,
            ["education", "formations", "diplômes", "academic background",
             "qualifications", "degrees", "schooling", "academic qualifications"]
        )
        education = self.edu_parser.parse(edu_section) if edu_section else []

        # 10. Certifications
        cert_section = self._find_section(
            sections,
            ["certifications", "attestations", "licenses", "credentials",
             "professional development", "training", "awards", "achievements"]
        )
        certifications = self.cert_ext.extract(cert_section) if cert_section else []
        if not certifications:
            certifications = self.cert_ext.extract(raw_text)

        # 11. Résumé
        summary_section = self._find_section(
            sections,
            ["summary", "résumé", "profil", "professional summary",
             "about", "about me", "profile", "objective", "career objective"]
        )

        # 12. Confiance (pour benchmarking)
        confidence = {
            "lang_detection": lang,
            "name":           "ner"          if ner_header["name"]       else "missing",
            "email":          "regex"         if email_m                  else "missing",
            "phone":          "regex"         if phone_m                  else "missing",
            "location":       "ner"           if ner_header["locations"]  else "missing",
            "skills":         "dict"          if skills                   else "missing",
            "languages":      ("dict+section" if lang_section else "dict_global") if languages else "missing",
            "experiences":    "heuristic"     if experiences              else "missing",
            "education":      "heuristic"     if education                else "missing",
            "certifications": "keyword"       if certifications           else "missing",
        }

        return CVData(
            full_name      = ner_header["name"],
            email          = email_m.group(0)              if email_m    else None,
            phone          = self._clean_phone(phone_m.group(0)) if phone_m else None,
            location       = ner_header["locations"][0]    if ner_header["locations"] else None,
            linkedin       = linkedin_m.group(0)           if linkedin_m else None,
            github         = github_m.group(0)             if github_m   else None,
            summary        = (summary_section or "")[:500] or None,
            skills         = skills,
            languages      = languages,
            experiences    = experiences,
            education      = education,
            certifications = certifications,
            detected_lang  = lang,
            _confidence    = confidence,
        )

    # ---------------------------------------------------------------
    def _find_section(self, sections: dict[str, str], keys: list[str]) -> Optional[str]:
        for key in keys:
            for section_key, content in sections.items():
                if key in section_key and content.strip():
                    return content
        return None

    @staticmethod
    def _clean_phone(raw: str) -> str:
        digits = re.sub(r"\D", "", raw)
        return raw.strip() if 7 <= len(digits) <= 15 else raw.strip()


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python cv_parser_classic_v3.py <chemin_du_cv> [--with-confidence] [--model MODEL]")
        sys.exit(1)

    forced = None
    if "--model" in sys.argv:
        idx    = sys.argv.index("--model")
        forced = sys.argv[idx + 1]

    parser = ClassicCVParser(spacy_model=forced)
    result = parser.parse(sys.argv[1])

    if "--with-confidence" in sys.argv:
        print(result.to_json_with_confidence())
    else:
        print(result.to_json())